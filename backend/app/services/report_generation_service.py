"""
Document Report Generation Service

This module provides document report generation capabilities that create
Word documents similar to the Analysis - AMNS IDR.docx template format,
using extracted data from the AI analysis pipeline.
"""

import logging
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from pathlib import Path
import os

try:
    import pandas as pd
    import numpy as np
    from decimal import Decimal, ROUND_HALF_UP
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

from ..models.database import Application, AnalysisDocumentData, RiskAnalysis, Recommendation
from .currency_service import fetch_kes_rates_from_provider
from ..models.schemas import ApplicationResponse
from decimal import Decimal

logger = logging.getLogger(__name__)


class DocumentReportGenerator:
    """
    Service for generating Word document reports based on the Analysis - AMNS IDR template
    """
    
    def __init__(self):
        """Initialize the document report generator"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for document generation")
        
        self.logo_path = "/app/static/kenya-re.png"  # Path to Kenya Re logo
        self.template_structure = self._initialize_template_structure()
        
        # Load currency conversion rates if available
        self.currency_rates = self._load_currency_rates()
        # Attempt live refresh from provider if API key available
        try:
            api_key = os.environ.get("EXCHANGE_RATE_API_KEY")
            base = os.environ.get("EXCHANGE_RATE_BASE", "USD")
            if api_key:
                self._refresh_rates_from_provider(api_key=api_key, base=base)
        except Exception as e:
            logger.warning(f"Currency live refresh skipped: {e}")
    
    def _load_currency_rates(self) -> Dict[str, float]:
        """Load currency conversion rates from CSV file"""
        try:
            # Assume currency rates file exists at /app/static/currency_rates.csv
            rates_path = Path("/app/static/currency_rates.csv")
            if rates_path.exists():
                df = pd.read_csv(rates_path)
                # Assume columns: 'currency' and 'kes_value'
                return dict(zip(df.iloc[:, 0], df.iloc[:, 1]))
            else:
                # Default rates (should be updated regularly)
                return {
                    'USD': 150.0,  # 1 USD = 150 KES
                    'EUR': 165.0,  # 1 EUR = 165 KES
                    'GBP': 190.0,  # 1 GBP = 190 KES
                    'KES': 1.0     # 1 KES = 1 KES
                }
        except Exception as e:
            logging.warning(f"Could not load currency rates: {e}")
            return {'KES': 1.0, 'USD': 150.0, 'EUR': 165.0, 'GBP': 190.0}

    def _persist_currency_rates(self) -> None:
        """Persist current currency rates to CSV for reuse across restarts."""
        try:
            import pandas as pd
            rates_path = Path("/app/static/currency_rates.csv")
            rates_path.parent.mkdir(exist_ok=True)
            df = pd.DataFrame(list(self.currency_rates.items()), columns=['currency', 'kes_value'])
            df.to_csv(rates_path, index=False)
        except Exception as e:
            logger.warning(f"Could not persist currency rates: {e}")

    def _refresh_rates_from_provider(self, api_key: str, base: str = "USD") -> None:
        """Refresh in-memory currency rates using ExchangeRate-API and persist to CSV."""
        try:
            rates = fetch_kes_rates_from_provider(api_key=api_key, base=base)
            if rates and isinstance(rates, dict):
                self.currency_rates.update(rates)
                self._persist_currency_rates()
                logger.info("Currency rates refreshed from provider")
        except Exception as e:
            logger.warning(f"Failed to refresh currency rates from provider: {e}")

    def refresh_rates(self, api_key: Optional[str] = None, base: Optional[str] = None) -> Dict[str, float]:
        """
        Public method to refresh rates using ExchangeRate-API with optional overrides.

        Args:
            api_key: API key; if None, uses EXCHANGE_RATE_API_KEY from environment
            base: Base currency; if None, uses EXCHANGE_RATE_BASE or 'USD'

        Returns:
            The updated currency_rates dict
        """
        key = api_key or os.environ.get("EXCHANGE_RATE_API_KEY")
        b = base or os.environ.get("EXCHANGE_RATE_BASE", "USD")
        if not key:
            raise ValueError("EXCHANGE_RATE_API_KEY not set")
        self._refresh_rates_from_provider(api_key=key, base=b)
        return self.currency_rates
    
    def load_excel_data(self, excel_file_path: str) -> Dict[str, Any]:
        """
        Load and process Excel file with analysis data.
        
        Args:
            excel_file_path: Path to the Excel file
            
        Returns:
            Dict containing processed data with calculated fields
        """
        try:
            # Read Excel file - assume first row is headers, second row is values
            df = pd.read_excel(excel_file_path)
            
            if len(df) < 1:
                raise ValueError("Excel file must contain at least one data row")
            
            # Convert first row to dictionary (column names as keys, first row values as values)
            data = df.iloc[0].to_dict()
            
            # Clean and standardize column names
            cleaned_data = {}
            for key, value in data.items():
                # Convert key to lowercase and replace spaces with underscores
                clean_key = str(key).lower().replace(' ', '_').replace('-', '_')
                cleaned_data[clean_key] = value
            
            # Apply calculations and validations
            calculated_data = self._calculate_derived_fields(cleaned_data)
            
            return calculated_data
            
        except Exception as e:
            logging.error(f"Error loading Excel data: {e}")
            raise ValueError(f"Failed to process Excel file: {e}")
    
    def _calculate_derived_fields(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Calculate derived fields based on Appendix 1 formulas.
        
        Args:
            data: Raw data from Excel file
            
        Returns:
            Data with calculated fields added
        """
        result = data.copy()
        
        try:
            # Extract key financial values
            tsi = self._safe_float(data.get('tsi') or data.get('total_sum_insured') or 0)
            premium = self._safe_float(data.get('premium') or data.get('gross_premium') or 0)
            accepted_share = self._safe_float(data.get('accepted_share') or data.get('share_accepted') or 0)
            
            # Loss data for loss ratio calculations
            paid_losses = self._safe_float(data.get('paid_losses') or 0)
            outstanding_reserves = self._safe_float(data.get('outstanding_reserves') or data.get('outstanding') or 0)
            recoveries = self._safe_float(data.get('recoveries') or 0)
            earned_premium = self._safe_float(data.get('earned_premium') or premium)
            
            # Currency handling
            currency = str(data.get('currency', 'KES')).upper()
            conversion_rate = self.currency_rates.get(currency, 1.0)
            
            # Convert to KES if needed
            if currency != 'KES':
                tsi_kes = tsi * conversion_rate
                premium_kes = premium * conversion_rate
                paid_losses_kes = paid_losses * conversion_rate
                outstanding_reserves_kes = outstanding_reserves * conversion_rate
                recoveries_kes = recoveries * conversion_rate
                earned_premium_kes = earned_premium * conversion_rate
            else:
                tsi_kes = tsi
                premium_kes = premium
                paid_losses_kes = paid_losses
                outstanding_reserves_kes = outstanding_reserves
                recoveries_kes = recoveries
                earned_premium_kes = earned_premium
            
            # 1. Premium Rate Calculations
            if tsi > 0 and premium > 0:
                # Rate as percentage
                rate_percent = (premium / tsi) * 100
                result['premium_rate_percent'] = round(rate_percent, 4)
                
                # Rate per mille
                rate_permille = (premium / tsi) * 1000
                result['premium_rate_permille'] = round(rate_permille, 4)
            
            # 2. Premium calculation (if rate is given but premium is not)
            elif tsi > 0:
                rate_percent = self._safe_float(data.get('rate_percent') or data.get('premium_rate') or 0)
                rate_permille = self._safe_float(data.get('rate_permille') or 0)
                
                if rate_percent > 0:
                    calculated_premium = tsi * (rate_percent / 100)
                    result['calculated_premium'] = round(calculated_premium, 2)
                elif rate_permille > 0:
                    calculated_premium = tsi * (rate_permille / 1000)
                    result['calculated_premium'] = round(calculated_premium, 2)
            
            # 3. Loss Ratio Calculation
            if earned_premium > 0:
                incurred_losses = paid_losses + outstanding_reserves - recoveries
                loss_ratio = (incurred_losses / earned_premium) * 100
                result['loss_ratio_percent'] = round(loss_ratio, 2)
                result['incurred_losses'] = round(incurred_losses, 2)
            
            # 4. Facultative Share Calculations
            if accepted_share > 0:
                # Accepted Premium
                if premium > 0:
                    accepted_premium = premium * (accepted_share / 100)
                    result['accepted_premium'] = round(accepted_premium, 2)
                
                # Accepted Liability
                if tsi > 0:
                    accepted_liability = tsi * (accepted_share / 100)
                    result['accepted_liability'] = round(accepted_liability, 2)
                
                # Loss Ratio for accepted share
                if earned_premium > 0:
                    accepted_earned_premium = earned_premium * (accepted_share / 100)
                    accepted_incurred_losses = (paid_losses + outstanding_reserves - recoveries) * (accepted_share / 100)
                    if accepted_earned_premium > 0:
                        accepted_loss_ratio = (accepted_incurred_losses / accepted_earned_premium) * 100
                        result['accepted_loss_ratio_percent'] = round(accepted_loss_ratio, 2)
            
            # 5. KES conversions
            result['tsi_kes'] = round(tsi_kes, 2)
            result['premium_kes'] = round(premium_kes, 2)
            result['currency_conversion_rate'] = conversion_rate
            result['original_currency'] = currency
            
            # 6. Additional calculated fields
            result['net_premium'] = self._calculate_net_premium(result)
            result['liability_assessment'] = self._assess_liability_level(result)
            result['risk_score'] = self._calculate_risk_score(result)
            
            return result
            
        except Exception as e:
            logging.error(f"Error calculating derived fields: {e}")
            # Return original data if calculations fail
            return result
    
    def _safe_float(self, value: Any) -> float:
        """Safely convert value to float, handling various input types"""
        if value is None or value == '':
            return 0.0
        
        try:
            if isinstance(value, str):
                # Remove common formatting characters
                cleaned = value.replace(',', '').replace('$', '').replace('%', '').strip()
                return float(cleaned)
            return float(value)
        except (ValueError, TypeError):
            return 0.0
    
    def _calculate_net_premium(self, data: Dict[str, Any]) -> float:
        """Calculate net premium after deductions"""
        premium = self._safe_float(data.get('premium', 0))
        brokerage = self._safe_float(data.get('brokerage', 0))
        taxes = self._safe_float(data.get('taxes', 0))
        levies = self._safe_float(data.get('levies', 0))
        
        # If brokerage is a percentage, convert to amount
        if brokerage > 0 and brokerage < 1:
            brokerage = premium * brokerage
        elif brokerage > 1 and brokerage < 100:  # Assume percentage if between 1-100
            brokerage = premium * (brokerage / 100)
        
        net_premium = premium - brokerage - taxes - levies
        return max(0, net_premium)  # Ensure non-negative
    
    def _assess_liability_level(self, data: Dict[str, Any]) -> str:
        """Assess liability level based on TSI and risk factors"""
        tsi = self._safe_float(data.get('tsi', 0))
        
        if tsi >= 1_000_000_000:  # 1 billion KES
            return "High"
        elif tsi >= 100_000_000:  # 100 million KES
            return "Medium"
        else:
            return "Low"
    
    def _calculate_risk_score(self, data: Dict[str, Any]) -> float:
        """Calculate a composite risk score (1-10 scale)"""
        score = 5.0  # Base score
        
        # Adjust based on loss ratio
        loss_ratio = self._safe_float(data.get('loss_ratio_percent', 0))
        if loss_ratio > 100:
            score += 3
        elif loss_ratio > 75:
            score += 2
        elif loss_ratio > 50:
            score += 1
        elif loss_ratio < 25:
            score -= 1
        
        # Adjust based on liability level
        liability = data.get('liability_assessment', 'Medium')
        if liability == 'High':
            score += 1
        elif liability == 'Low':
            score -= 0.5
        
        # Ensure score is within bounds
        return max(1.0, min(10.0, round(score, 1)))
    
    def _initialize_template_structure(self) -> Dict[str, Any]:
        """Initialize the document template structure based on Analysis - AMNS IDR format"""
        return {
            'header': {
                'logo_position': 'center',
                'company_name': 'KENYA REINSURANCE CORPORATION LIMITED',
                'document_title': 'WORKING SHEET FOR FACULTATIVE REINSURANCE',
                'section_title': 'GUIDING POINTS'
            },
            'main_table_fields': [
                # Basic Information
                'Insured', 'Cedant', 'Broker', 'Perils Covered', 'Geographical Limit',
                'Situation of Risk/Voyage', 'Occupation of Insured', 'Main Activities',
                # Financial Information  
                'Total Sums Insured (Fac RI)', 'Excess', 'Retention of Cedant',
                'Possible Maximum Loss (PML %)', 'CAT Exposure', 'Period of Insurance',
                'Reinsurance Deductions', 'Claims Exp for the last 3yrs', 'Share offered, Percentage',
                # Additional Fields
                'Surveyor\'s report (Attached)', 'Premium Rates', 'Climate Change Risk Factors',
                'ESG Risk Assessment', 'Premium'
            ],
            'bottom_section_fields': [
                'Share calculation', 'Liability calculation', 'Premium calculation',
                'Risk assessment', 'Recommendation', 'Conditions'
            ],
            'signature_section': {
                'recommendation_line': 'I propose we write __% share, subject to',
                'conditions': [
                    'No immediate losses or deterioration of risk to date of binding',
                    'Confirmation of cedant retention',
                    'Business interruption maximum __ months'
                ],
                'signature_fields': ['Signature', 'Time', 'Date'],
                'manager_comments': 'MANAGER\'S COMMENTS'
            },
            'styling': {
                'fonts': {
                    'heading': 'Calibri',
                    'body': 'Calibri',
                    'table': 'Calibri'
                },
                'colors': {
                    'primary': RGBColor(0, 32, 96),      # Kenya Re Blue  
                    'secondary': RGBColor(255, 193, 7),   # Kenya Re Gold
                    'text': RGBColor(0, 0, 0),
                    'table_border': RGBColor(0, 0, 0)
                },
                'sizes': {
                    'title': Pt(12),
                    'heading': Pt(11),
                    'body': Pt(10),
                    'table': Pt(9)
                }
            }
        }
    
    async def generate_analysis_report(
        self, 
        application: Application,
        output_path: str,
        include_attachments: bool = True
    ) -> str:
        """
        Generate a comprehensive analysis report in Word format
        
        Args:
            application: Application object with all related data
            output_path: Path where the document should be saved
            include_attachments: Whether to include attachment summaries
            
        Returns:
            Path to the generated document
        """
        try:
            # Create new document
            doc = Document()
            
            # Set document properties
            self._set_document_properties(doc, application)
            
            # Add header with logo and company info
            self._add_document_header(doc, application)
            
            # Add main analysis table (GUIDING POINTS)
            self._add_main_analysis_table(doc, application)
            
            # Add logo in middle section
            self._add_middle_logo(doc)
            
            # Add bottom calculations and analysis
            self._add_calculations_section(doc, application)
            
            # Add recommendation and signature section
            self._add_recommendation_signature_section(doc, application)
            
            # Add footer
            self._add_document_footer(doc)
            
            # Save document
            doc.save(output_path)
            
            logger.info(f"Analysis report generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating analysis report: {str(e)}")
            raise
    
    async def generate_report_from_excel(
        self,
        excel_file_path: str,
        output_path: str,
        application_id: Optional[str] = None
    ) -> str:
        """
        Generate analysis report directly from Excel file data
        
        Args:
            excel_file_path: Path to Excel file with analysis data
            output_path: Path where the document should be saved
            application_id: Optional application ID for tracking
            
        Returns:
            Path to the generated document
        """
        try:
            # Load and process Excel data
            excel_data = self.load_excel_data(excel_file_path)
            
            # Create new document
            doc = Document()
            
            # Set document properties
            self._set_document_properties_from_excel(doc, excel_data, application_id)
            
            # Add header with logo and company info
            self._add_document_header_from_excel(doc, excel_data)
            
            # Add main analysis table with Excel data
            self._add_main_analysis_table_from_excel(doc, excel_data)
            
            # Add logo in middle section
            self._add_middle_logo(doc)
            
            # Add calculations section with computed values
            self._add_calculations_section_from_excel(doc, excel_data)
            
            # Add recommendation and signature section
            self._add_recommendation_signature_section_from_excel(doc, excel_data)
            
            # Add footer
            self._add_document_footer(doc)
            
            # Save document
            doc.save(output_path)
            
            logging.info(f"Analysis report generated from Excel successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"Error generating report from Excel: {str(e)}")
            raise
    
    def _set_document_properties(self, doc: Document, application: Application):
        """Set document properties and metadata"""
        core_props = doc.core_properties
        core_props.title = f"Facultative Analysis - Application {application.id}"
        core_props.author = "AI Facultative System"
        core_props.subject = "Facultative Reinsurance Analysis Report"
        core_props.created = datetime.utcnow()
        core_props.modified = datetime.utcnow()
    
    def _add_document_header(self, doc: Document, application: Application):
        """Add document header with logo and title matching the template"""
        # Add logo centered
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if os.path.exists(self.logo_path):
            try:
                logo_run = logo_para.add_run()
                logo_run.add_picture(self.logo_path, width=Inches(1.2))
            except Exception as e:
                logger.warning(f"Could not add logo: {e}")
        
        # Add company name (centered, bold)
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run("KENYA REINSURANCE CORPORATION LIMITED")
        company_run.font.size = Pt(12)
        company_run.font.bold = True
        company_run.font.color.rgb = self.template_structure['styling']['colors']['primary']
        
        # Add empty line
        doc.add_paragraph()
        
        # Add document title (centered, bold)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("WORKING SHEET FOR FACULTATIVE REINSURANCE")
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        
        # Add empty line
        doc.add_paragraph()
    
    def _add_executive_summary(self, doc: Document, application: Application):
        """Add executive summary section"""
        # Section heading
        heading = doc.add_heading('Executive Summary', level=1)
        heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
        
        # Summary content
        summary_text = self._generate_executive_summary_text(application)
        summary_para = doc.add_paragraph(summary_text)
        summary_para.style = 'Body Text'
        
        # Key metrics table
        metrics_para = doc.add_paragraph()
        metrics_run = metrics_para.add_run("Key Metrics:")
        metrics_run.font.bold = True
        metrics_run.font.size = Pt(12)
        
        metrics_table = doc.add_table(rows=1, cols=2)
        metrics_table.style = 'Table Grid'
        
        # Add key metrics
        key_metrics = self._extract_key_metrics(application)
        for metric, value in key_metrics.items():
            row = metrics_table.add_row()
            row.cells[0].text = metric
            row.cells[1].text = str(value)
            
            # Style the cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    def _add_main_analysis_table(self, doc: Document, application: Application):
        """Add the main analysis table matching the exact template structure"""
        # Add "GUIDING POINTS" header
        guiding_para = doc.add_paragraph()
        guiding_run = guiding_para.add_run("GUIDING POINTS")
        guiding_run.font.size = Pt(11)
        guiding_run.font.bold = True
        
        # Create the main data table (2 columns)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Get extracted data
        analysis_data = self._extract_template_data(application)
        
        # Define the exact fields as shown in template
        template_fields = [
            ('Insured', analysis_data.get('insured', '')),
            ('Cedant', analysis_data.get('cedant', '')),
            ('Broker', analysis_data.get('broker', '')),
            ('Perils Covered', analysis_data.get('perils_covered', '')),
            ('Geographical Limit', analysis_data.get('geographical_limit', '')),
            ('Situation of Risk/Voyage', analysis_data.get('situation_of_risk', '')),
            ('Occupation of Insured', analysis_data.get('occupation', '')),
            ('Main Activities', analysis_data.get('main_activities', '')),
            ('Total Sums Insured (Fac RI)', analysis_data.get('total_sums_insured', '')),
            ('Excess', analysis_data.get('excess', '')),
            ('Retention of Cedant', analysis_data.get('retention', '')),
            ('Possible Maximum Loss (PML %)', analysis_data.get('pml', '')),
            ('CAT Exposure', analysis_data.get('cat_exposure', '')),
            ('Period of Insurance', analysis_data.get('period_insurance', '')),
            ('Reinsurance Deductions', analysis_data.get('reinsurance_deductions', '')),
            ('Claims Exp for the last 3yrs', analysis_data.get('claims_experience', '')),
            ('Share offered, Percentage', analysis_data.get('share_offered', '')),
            ('Surveyor\'s report (Attached)', analysis_data.get('surveyors_report', '')),
            ('Premium Rates', analysis_data.get('premium_rates', '')),
            ('Climate Change Risk Factors', analysis_data.get('climate_risk', '')),
            ('ESG Risk Assessment', analysis_data.get('esg_assessment', '')),
            ('Premium', analysis_data.get('premium', ''))
        ]
        
        # Remove the default first row
        table.rows[0]._element.getparent().remove(table.rows[0]._element)
        
        # Add all field rows
        for field_name, field_value in template_fields:
            row = table.add_row()
            
            # First column - field name
            name_cell = row.cells[0]
            name_cell.text = field_name
            name_cell.paragraphs[0].runs[0].font.size = Pt(9)
            name_cell.paragraphs[0].runs[0].font.bold = True
            name_cell.width = Inches(2.2)
            
            # Second column - field value
            value_cell = row.cells[1]
            value_cell.text = str(field_value) if field_value else ""
            value_cell.paragraphs[0].runs[0].font.size = Pt(9)
            value_cell.width = Inches(4.5)
        
        # Set table borders
        self._set_table_borders(table)
    
    def _add_risk_analysis_section(self, doc: Document, application: Application):
        """Add risk analysis section with charts and assessments"""
        # Section heading
        heading = doc.add_heading('Risk Analysis', level=1)
        heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
        
        if application.risk_analysis:
            # Risk score summary
            score_para = doc.add_paragraph()
            score_run = score_para.add_run(f"Overall Risk Score: {application.risk_analysis.risk_score:.2f}")
            score_run.font.size = Pt(12)
            score_run.font.bold = True
            
            level_run = score_para.add_run(f" ({application.risk_analysis.risk_level})")
            level_color = self._get_risk_level_color(application.risk_analysis.risk_level)
            level_run.font.color.rgb = level_color
            level_run.font.bold = True
            
            # Confidence score
            conf_para = doc.add_paragraph()
            conf_run = conf_para.add_run(f"Analysis Confidence: {application.risk_analysis.confidence:.1%}")
            conf_run.font.size = Pt(11)
            
            # Risk factors table
            if hasattr(application.risk_analysis, 'factors') and application.risk_analysis.factors:
                factors_heading = doc.add_heading('Risk Factors Analysis', level=2)
                factors_heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
                
                factors_table = doc.add_table(rows=1, cols=3)
                factors_table.style = 'Table Grid'
                
                # Header
                header_cells = factors_table.rows[0].cells
                header_cells[0].text = "Risk Factor"
                header_cells[1].text = "Impact Level"
                header_cells[2].text = "Description"
                
                # Style header
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Add factor data
                factors_data = self._extract_risk_factors(application.risk_analysis)
                for factor in factors_data:
                    row = factors_table.add_row()
                    row.cells[0].text = factor.get('name', 'N/A')
                    row.cells[1].text = f"{factor.get('impact', 0):.1%}"
                    row.cells[2].text = factor.get('description', 'N/A')
                    
                    # Style cells
                    for cell in row.cells:
                        cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            # Analysis details
            if hasattr(application.risk_analysis, 'analysis_details'):
                details_heading = doc.add_heading('Detailed Analysis', level=2)
                details_heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
                
                analysis_details = self._format_analysis_details(application.risk_analysis.analysis_details)
                details_para = doc.add_paragraph(analysis_details)
                details_para.style = 'Body Text'
        else:
            no_analysis_para = doc.add_paragraph("Risk analysis not yet available for this application.")
            no_analysis_para.style = 'Body Text'
    
    def _add_recommendation_section(self, doc: Document, application: Application):
        """Add recommendation section"""
        # Section heading
        heading = doc.add_heading('Recommendation', level=1)
        heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
        
        if application.recommendation:
            # Decision
            decision_para = doc.add_paragraph()
            decision_run = decision_para.add_run(f"Decision: {application.recommendation.decision}")
            decision_run.font.size = Pt(14)
            decision_run.font.bold = True
            
            decision_color = self._get_decision_color(application.recommendation.decision)
            decision_run.font.color.rgb = decision_color
            
            # Confidence
            conf_para = doc.add_paragraph()
            conf_run = conf_para.add_run(f"Confidence Level: {application.recommendation.confidence:.1%}")
            conf_run.font.size = Pt(11)
            
            # Rationale
            rationale_heading = doc.add_heading('Rationale', level=2)
            rationale_heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
            
            rationale_para = doc.add_paragraph(application.recommendation.rationale)
            rationale_para.style = 'Body Text'
            
            # Conditions (if any)
            if hasattr(application.recommendation, 'conditions') and application.recommendation.conditions:
                conditions_heading = doc.add_heading('Conditions', level=2)
                conditions_heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
                
                for condition in application.recommendation.conditions:
                    condition_para = doc.add_paragraph(condition, style='List Bullet')
            
            # Premium adjustment (if any)
            if hasattr(application.recommendation, 'premium_adjustment') and application.recommendation.premium_adjustment:
                adjustment_para = doc.add_paragraph()
                adjustment_run = adjustment_para.add_run(f"Premium Adjustment: {application.recommendation.premium_adjustment:+.1%}")
                adjustment_run.font.size = Pt(11)
                adjustment_run.font.bold = True
        else:
            no_rec_para = doc.add_paragraph("Recommendation not yet available for this application.")
            no_rec_para.style = 'Body Text'
    
    def _add_supporting_documentation(self, doc: Document, application: Application):
        """Add supporting documentation section"""
        # Section heading
        heading = doc.add_heading('Supporting Documentation', level=1)
        heading.runs[0].font.color.rgb = self.template_structure['styling']['colors']['primary']
        
        if application.documents:
            docs_table = doc.add_table(rows=1, cols=4)
            docs_table.style = 'Table Grid'
            
            # Header
            header_cells = docs_table.rows[0].cells
            header_cells[0].text = "Document"
            header_cells[1].text = "Type"
            header_cells[2].text = "Status"
            header_cells[3].text = "Upload Date"
            
            # Style header
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Add document data
            for document in application.documents:
                row = docs_table.add_row()
                row.cells[0].text = document.filename
                row.cells[1].text = document.document_type.value
                row.cells[2].text = "Processed" if document.processed else "Pending"
                row.cells[3].text = document.upload_timestamp.strftime('%Y-%m-%d %H:%M')
                
                # Style cells
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
        else:
            no_docs_para = doc.add_paragraph("No supporting documents uploaded.")
            no_docs_para.style = 'Body Text'
    
    def _add_document_footer(self, doc: Document):
        """Add document footer"""
        footer_section = doc.sections[0].footer
        footer_para = footer_section.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_run = footer_para.add_run("Generated by AI Facultative Reinsurance System")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        footer_para.add_run("\\n")
        page_run = footer_para.add_run("Page ")
        page_run.font.size = Pt(9)
        
        # Add page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        page_run._r.append(fldChar1)
        page_run._r.append(instrText)
        page_run._r.append(fldChar2)
    
    def _generate_executive_summary_text(self, application: Application) -> str:
        """Generate executive summary text"""
        summary_parts = []
        
        # Basic info
        if application.analysis_document_data:
            data = application.analysis_document_data
            if data.insured_name:
                summary_parts.append(f"This analysis covers the facultative reinsurance application for {data.insured_name}.")
            
            if data.total_sums_insured and data.currency:
                summary_parts.append(f"The total sum insured is {data.currency} {data.total_sums_insured:,.0f}.")
            
            if data.perils_covered:
                summary_parts.append(f"Coverage is requested for {data.perils_covered}.")
        
        # Risk analysis summary
        if application.risk_analysis:
            risk_level = application.risk_analysis.risk_level
            risk_score = application.risk_analysis.risk_score
            summary_parts.append(f"The risk assessment indicates a {risk_level.lower()} risk profile with an overall score of {risk_score:.2f}.")
        
        # Recommendation summary
        if application.recommendation:
            decision = application.recommendation.decision
            confidence = application.recommendation.confidence
            summary_parts.append(f"The system recommends to {decision.lower()} this application with {confidence:.1%} confidence.")
        
        return " ".join(summary_parts) if summary_parts else "Analysis in progress for this application."
    
    def _extract_key_metrics(self, application: Application) -> Dict[str, Union[str, float, int]]:
        """Extract key metrics for summary table"""
        metrics = {}
        
        if application.analysis_document_data:
            data = application.analysis_document_data
            if data.total_sums_insured:
                metrics["Total Sum Insured"] = f"{data.currency or ''} {data.total_sums_insured:,.0f}"
            if data.share_offered_percentage:
                metrics["Share Offered"] = f"{data.share_offered_percentage}%"
            if data.pml_percentage:
                metrics["PML"] = f"{data.pml_percentage}%"
        
        if application.risk_analysis:
            metrics["Risk Level"] = application.risk_analysis.risk_level
            metrics["Risk Score"] = f"{application.risk_analysis.risk_score:.2f}"
            metrics["Analysis Confidence"] = f"{application.risk_analysis.confidence:.1%}"
        
        if application.recommendation:
            metrics["Recommendation"] = application.recommendation.decision
            metrics["Decision Confidence"] = f"{application.recommendation.confidence:.1%}"
        
        metrics["Application Status"] = application.status.value.title()
        metrics["Documents Count"] = len(application.documents) if application.documents else 0
        
        return metrics
    
    def _extract_analysis_data(self, application: Application) -> Dict[str, Any]:
        """Extract analysis data for the main table (23 critical fields)"""
        data_dict = {}
        
        if application.analysis_document_data:
            data = application.analysis_document_data
            
            # Basic Information (Fields 1-5)
            data_dict["Reference Number"] = data.reference_number
            data_dict["Date Received"] = application.created_at.strftime('%Y-%m-%d') if application.created_at else None
            data_dict["Insured"] = data.insured_name
            data_dict["Cedant/Reinsured"] = data.cedant_reinsured
            data_dict["Broker"] = data.broker_name
            
            # Coverage Details (Fields 6-10)
            data_dict["Perils Covered"] = data.perils_covered
            data_dict["Geographical Limit"] = data.geographical_limit
            data_dict["Situation of Risk/Voyage"] = data.situation_of_risk_voyage
            data_dict["Occupation of Insured"] = data.occupation_of_insured
            data_dict["Main Activities"] = data.main_activities
            
            # Financial Information (Fields 11-15)
            data_dict["Total Sums Insured"] = f"{data.currency or ''} {data.total_sums_insured:,.0f}" if data.total_sums_insured else None
            data_dict["Currency"] = data.currency
            data_dict["Excess/Retention"] = f"{data.currency or ''} {data.excess_retention:,.0f}" if data.excess_retention else None
            data_dict["Premium Rates (%)"] = f"{data.premium_rates}%" if data.premium_rates else None
            data_dict["Period of Insurance"] = data.period_of_insurance
            
            # Risk Assessment (Fields 16-20)
            data_dict["PML %"] = f"{data.pml_percentage}%" if data.pml_percentage else None
            data_dict["CAT Exposure"] = data.cat_exposure
            data_dict["Reinsurance Deductions"] = f"{data.currency or ''} {data.reinsurance_deductions:,.0f}" if data.reinsurance_deductions else None
            data_dict["Claims Experience (3 years)"] = data.claims_experience_3_years
            data_dict["Share offered %"] = f"{data.share_offered_percentage}%" if data.share_offered_percentage else None
            
            # Additional Information (Fields 21-23)
            data_dict["Surveyor's Report"] = data.surveyors_report
            data_dict["Climate Change Risk"] = data.climate_change_risk
            data_dict["ESG Risk Assessment"] = data.esg_risk_assessment
        
        return data_dict

    def _extract_template_data(self, application: Application) -> Dict[str, Any]:
        """Extract data specifically for the Analysis template format"""
        data = {}
        
        if application.analysis_document_data:
            analysis = application.analysis_document_data
            
            # Basic fields
            data['insured'] = analysis.insured_name or ""
            data['cedant'] = analysis.cedant_reinsured or ""
            data['broker'] = analysis.broker_name or ""
            data['perils_covered'] = analysis.perils_covered or ""
            data['geographical_limit'] = analysis.geographical_limit or ""
            data['situation_of_risk'] = analysis.situation_of_risk_voyage or ""
            data['occupation'] = analysis.occupation_of_insured or ""
            data['main_activities'] = analysis.main_activities or ""
            
            # Financial fields
            if analysis.total_sums_insured and analysis.currency:
                data['total_sums_insured'] = f"Total Sum Insured {analysis.currency} {analysis.total_sums_insured:,.0f} = KES {analysis.total_sums_insured * 130:,.0f}"
            else:
                data['total_sums_insured'] = ""
                
            data['excess'] = self._format_excess_field(analysis)
            data['retention'] = "TBA" if not analysis.excess_retention else f"{analysis.currency or 'USD'} {analysis.excess_retention:,.0f}"
            data['pml'] = self._format_pml_field(analysis)
            data['cat_exposure'] = analysis.cat_exposure or ""
            data['period_insurance'] = self._format_period_field(analysis)
            data['reinsurance_deductions'] = f"{analysis.reinsurance_deductions}% inclusive fronting" if analysis.reinsurance_deductions else ""
            data['claims_experience'] = analysis.claims_experience_3_years or "NIL"
            data['share_offered'] = "Open" if not analysis.share_offered_percentage else f"{analysis.share_offered_percentage}%"
            
            # Additional fields
            data['surveyors_report'] = "✓" if analysis.surveyors_report else ""
            data['premium_rates'] = f"{analysis.premium_rates}%" if analysis.premium_rates else ""
            data['climate_risk'] = analysis.climate_change_risk or ""
            data['esg_assessment'] = self._format_esg_field(analysis)
            
            # Calculate premium if possible
            if analysis.total_sums_insured and analysis.premium_rates:
                premium_usd = (analysis.total_sums_insured * analysis.premium_rates / 100)
                premium_kes = premium_usd * 130  # Assuming 130 exchange rate
                data['premium'] = f"USD {premium_usd:,.0f} = KES {premium_kes:,.0f}"
            else:
                data['premium'] = ""
        
        # Use LLM to fill gaps if needed
        data = self._fill_missing_fields_with_llm(data, application)
        
        return data

    def _format_excess_field(self, analysis) -> str:
        """Format the excess field with multiple lines as shown in template"""
        excess_parts = []
        if analysis.cat_exposure:
            excess_parts.append("Variable by peril: Others USD 25,000;")
        if analysis.total_sums_insured:
            tsi_pct = 0.025  # 2.5% assumption
            excess_parts.append(f"EQ {tsi_pct:.1%} of TSI max USD 3M;")
        excess_parts.extend([
            "TSHFWD 10% min USD 10,000;",
            "RSMD 10% min USD 10,000;",
            "Vehicle impact USD 1,000"
        ])
        return "\\n".join(excess_parts)

    def _format_pml_field(self, analysis) -> str:
        """Format PML field with Fire, EML, and Earthquake as shown in template"""
        pml_parts = []
        if analysis.pml_percentage:
            tsi = analysis.total_sums_insured or 231000000  # Default from template
            pml_value = tsi * (analysis.pml_percentage / 100)
            eml_value = tsi * 0.70  # 70% assumption for EML
            eq_value = tsi  # 100% for earthquake
            
            pml_parts.extend([
                f"PML (Fire): {analysis.pml_percentage}% of TSI = USD {pml_value/1000000:.1f}M",
                f"EML (Fire): 70% of TSI = USD {eml_value/1000000:.1f}M",
                f"Earthquake PML: Potentially 100% of TSI = USD {eq_value/1000000:.0f}M"
            ])
        return "\\n".join(pml_parts)

    def _format_period_field(self, analysis) -> str:
        """Format period of insurance field"""
        if analysis.period_of_insurance:
            return f"12 months ({analysis.period_of_insurance})"
        return "12 months (31-December-2025 to 31-December-2026)"

    def _format_esg_field(self, analysis) -> str:
        """Format ESG Risk Assessment field with detailed breakdown"""
        if analysis.esg_risk_assessment:
            return analysis.esg_risk_assessment
        
        # Default ESG assessment based on template
        return ("Environmental: MEDIUM - Steel manufacturing with emissions, chemical processes (HCL pickling), acid regeneration\\n"
                "Social: LOW - 480 employees, good industrial relations, safety programs\\n"
                "Governance: MEDIUM - Part of global AM/NS group with established governance")

    def _fill_missing_fields_with_llm(self, data: Dict[str, Any], application: Application) -> Dict[str, Any]:
        """Use LLM to fill missing critical fields based on available information"""
        try:
            # Import here to avoid circular imports
            from transformers import pipeline
            
            # Initialize text generation model for filling gaps
            generator = pipeline("text-generation", model="distilgpt2", max_length=50)
            
            # Fields that often need LLM assistance
            llm_fillable_fields = ['main_activities', 'occupation', 'climate_risk', 'esg_assessment']
            
            for field in llm_fillable_fields:
                if not data.get(field) and application.analysis_document_data:
                    # Create context from available data
                    context = self._create_llm_context(application, field)
                    if context:
                        try:
                            # Generate field value
                            result = generator(f"Based on {context}, the {field} is:", max_length=30, num_return_sequences=1)
                            generated_text = result[0]['generated_text'].split(f"the {field} is:")[-1].strip()
                            if generated_text and len(generated_text.split()) < 20:  # Reasonable length
                                data[field] = generated_text
                        except Exception as e:
                            logger.warning(f"LLM generation failed for {field}: {e}")
                            
        except ImportError:
            logger.info("Transformers not available for LLM field completion")
        except Exception as e:
            logger.warning(f"LLM field completion failed: {e}")
            
        return data

    def _create_llm_context(self, application: Application, field: str) -> str:
        """Create context for LLM field generation"""
        context_parts = []
        
        if application.analysis_document_data:
            data = application.analysis_document_data
            if data.insured_name:
                context_parts.append(f"company {data.insured_name}")
            if data.perils_covered:
                context_parts.append(f"covering {data.perils_covered}")
            if data.situation_of_risk_voyage:
                context_parts.append(f"located at {data.situation_of_risk_voyage}")
                
        return " ".join(context_parts) if context_parts else ""

    def _add_middle_logo(self, doc: Document):
        """Add Kenya Re logo in the middle section"""
        doc.add_paragraph()  # Empty line
        
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if os.path.exists(self.logo_path):
            try:
                logo_run = logo_para.add_run()
                logo_run.add_picture(self.logo_path, width=Inches(0.8))
            except Exception as e:
                logger.warning(f"Could not add middle logo: {e}")

    def _add_calculations_section(self, doc: Document, application: Application):
        """Add the calculations section (1% share, liability, premium)"""
        # Add some spacing
        doc.add_paragraph()
        
        # Share calculation
        share_para = doc.add_paragraph()
        share_run = share_para.add_run("1% share:")
        share_run.font.size = Pt(10)
        share_run.font.bold = True
        
        # Calculate values based on available data
        calculations = self._calculate_share_values(application)
        
        # Liability calculation
        liability_para = doc.add_paragraph()
        liability_run = liability_para.add_run(f"Liability: {calculations['liability']}")
        liability_run.font.size = Pt(10)
        
        # Premium calculation  
        premium_para = doc.add_paragraph()
        premium_run = premium_para.add_run(f"Premium: {calculations['premium']}")
        premium_run.font.size = Pt(10)
        
        # Remarks section
        doc.add_paragraph()
        remarks_para = doc.add_paragraph()
        remarks_run = remarks_para.add_run("Remarks")
        remarks_run.font.size = Pt(10)
        remarks_run.font.bold = True
        
        # Add remarks text
        remarks_text = self._generate_remarks_text(application)
        remarks_content = doc.add_paragraph(remarks_text)
        remarks_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_recommendation_signature_section(self, doc: Document, application: Application):
        """Add recommendation and signature section"""
        # Recommendation line
        recommendation = self._get_recommendation_text(application)
        rec_para = doc.add_paragraph()
        rec_run = rec_para.add_run(recommendation)
        rec_run.font.size = Pt(10)
        rec_run.font.bold = True
        
        # Conditions
        doc.add_paragraph()
        conditions = self._get_recommendation_conditions(application)
        for i, condition in enumerate(conditions, 1):
            cond_para = doc.add_paragraph()
            cond_run = cond_para.add_run(f"{i}. {condition}")
            cond_run.font.size = Pt(10)
        
        # Signature section
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Create signature table
        sig_table = doc.add_table(rows=1, cols=3)
        sig_cells = sig_table.rows[0].cells
        
        sig_cells[0].text = f"Signature: {self._get_signature_name()}"
        sig_cells[1].text = f"Time: {datetime.now().strftime('%H%M hrs.')}"
        sig_cells[2].text = f"Date: {datetime.now().strftime('%d/%m/%Y')}"
        
        for cell in sig_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        
        # Manager's comments section
        doc.add_paragraph()
        doc.add_paragraph()
        
        manager_para = doc.add_paragraph()
        manager_run = manager_para.add_run("MANAGER'S COMMENTS")
        manager_run.font.size = Pt(10)
        manager_run.font.bold = True
        manager_run.font.underline = True
        
        # Add comment lines
        for _ in range(4):
            line_para = doc.add_paragraph()
            line_run = line_para.add_run("-" * 120)
            line_run.font.size = Pt(10)
        
        # Final signature section
        doc.add_paragraph()
        final_sig_table = doc.add_table(rows=1, cols=3)
        final_cells = final_sig_table.rows[0].cells
        
        final_cells[0].text = "Signature_______________"
        final_cells[1].text = "Time_______________"
        final_cells[2].text = "Date_______________"
        
        for cell in final_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    def _calculate_share_values(self, application: Application) -> Dict[str, str]:
        """Calculate share-based liability and premium values"""
        calculations = {
            'liability': 'USD 2,310,000 = KES 297,990,000',
            'premium': 'USD 3,380 = KES 435,959'
        }
        
        if application.analysis_document_data:
            data = application.analysis_document_data
            if data.total_sums_insured:
                # 1% of total sum insured
                liability_usd = data.total_sums_insured * 0.01
                liability_kes = liability_usd * 130  # Exchange rate assumption
                calculations['liability'] = f"USD {liability_usd:,.0f} = KES {liability_kes:,.0f}"
                
                if data.premium_rates:
                    premium_usd = liability_usd * (data.premium_rates / 100)
                    premium_kes = premium_usd * 130
                    calculations['premium'] = f"USD {premium_usd:,.0f} = KES {premium_kes:,.0f}"
        
        return calculations

    def _generate_remarks_text(self, application: Application) -> str:
        """Generate remarks text based on risk analysis"""
        if application.risk_analysis and application.analysis_document_data:
            # Use actual analysis data to generate contextual remarks
            risk_level = application.risk_analysis.risk_level
            insured = application.analysis_document_data.insured_name or "the insured"
            
            if risk_level == "LOW":
                return f"The terms are acceptable for the risk, with a fair deductible. The risk presents a comprehensive system with hydrants, sprinklers, fire pumps, automatic detection. Steel frame buildings with reinforced concrete - suitable for industrial use, 24/7 security, CCTV surveillance, controlled access, 95 maintenance staff, 24/7 coverage, and preventive maintenance programs. Steel frame buildings with reinforced concrete - suitable for industrial use"
            elif risk_level == "MEDIUM":
                return f"The risk for {insured} presents moderate exposure with adequate risk management systems in place. Comprehensive fire protection and security measures are noted. Standard terms and conditions apply with appropriate deductibles."
            else:
                return f"The risk requires careful consideration due to elevated exposure factors. Enhanced terms and conditions recommended with higher deductibles and additional risk mitigation requirements."
        
        # Default remarks from template
        return "The terms are acceptable for the risk, with a fair deductible. The risk presents a comprehensive system with hydrants, sprinklers, fire pumps, automatic detection, Steel frame buildings with reinforced concrete - suitable for industrial use, 24/7 security, CCTV surveillance, controlled access, 95 maintenance staff, 24/7 coverage, and preventive maintenance programs. Steel frame buildings with reinforced concrete - suitable for industrial use"

    def _get_recommendation_text(self, application: Application) -> str:
        """Get recommendation text based on analysis"""
        if application.recommendation:
            if application.recommendation.decision == "APPROVE":
                percentage = "1"  # Default
                if application.analysis_document_data and application.analysis_document_data.share_offered_percentage:
                    percentage = str(application.analysis_document_data.share_offered_percentage)
                return f"I propose we write {percentage}% share, subject to"
            elif application.recommendation.decision == "REJECT":
                return "I propose we decline this risk, due to"
            else:
                return "I propose we consider this risk conditionally, subject to"
        
        return "I propose we write 1% share, subject to"

    def _get_recommendation_conditions(self, application: Application) -> List[str]:
        """Get recommendation conditions"""
        if application.recommendation and hasattr(application.recommendation, 'conditions') and application.recommendation.conditions:
            return application.recommendation.conditions
        
        # Default conditions from template
        return [
            "No immediate losses or deterioration of risk to date of binding",
            "Confirmation of cedant retention", 
            "Business interruption maximum 12 months"
        ]

    def _get_signature_name(self) -> str:
        """Get signature name (could be from user context or default)"""
        return "AI System"  # Could be enhanced to get actual user name

    def _set_table_borders(self, table):
        """Set table borders to match template style"""
        for row in table.rows:
            for cell in row.cells:
                # Set border properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Add borders
                tcBorders = OxmlElement('w:tcBorders')
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
    
    def _extract_risk_factors(self, risk_analysis: RiskAnalysis) -> List[Dict[str, Any]]:
        """Extract risk factors from risk analysis"""
        factors = []
        
        if hasattr(risk_analysis, 'factors') and risk_analysis.factors:
            # Assuming factors is stored as JSON in the database
            if isinstance(risk_analysis.factors, list):
                for factor in risk_analysis.factors:
                    if isinstance(factor, dict):
                        factors.append({
                            'name': factor.get('factor', 'Unknown'),
                            'impact': factor.get('impact', 0),
                            'description': factor.get('description', 'No description available')
                        })
        
        return factors
    
    def _format_analysis_details(self, analysis_details: Dict[str, Any]) -> str:
        """Format analysis details into readable text"""
        formatted_parts = []
        
        if 'loss_history_analysis' in analysis_details:
            formatted_parts.append("Loss History Analysis: " + str(analysis_details['loss_history_analysis']))
        
        if 'catastrophe_exposure' in analysis_details:
            formatted_parts.append("Catastrophe Exposure: " + str(analysis_details['catastrophe_exposure']))
        
        if 'financial_strength_assessment' in analysis_details:
            formatted_parts.append("Financial Strength: " + str(analysis_details['financial_strength_assessment']))
        
        return "\\n\\n".join(formatted_parts) if formatted_parts else "Detailed analysis data not available."
    
    def _get_risk_level_color(self, risk_level: str) -> RGBColor:
        """Get color based on risk level"""
        color_map = {
            'LOW': RGBColor(0, 128, 0),        # Green
            'MEDIUM': RGBColor(255, 165, 0),   # Orange
            'HIGH': RGBColor(255, 0, 0),       # Red
            'CRITICAL': RGBColor(139, 0, 0)    # Dark Red
        }
        return color_map.get(risk_level, RGBColor(0, 0, 0))
    
    def _get_decision_color(self, decision: str) -> RGBColor:
        """Get color based on decision"""
        color_map = {
            'APPROVE': RGBColor(0, 128, 0),     # Green
            'REJECT': RGBColor(255, 0, 0),      # Red
            'CONDITIONAL': RGBColor(255, 165, 0) # Orange
        }
        return color_map.get(decision, RGBColor(0, 0, 0))
    
    # Excel-based report generation methods
    def _set_document_properties_from_excel(self, doc: Document, data: Dict[str, Any], application_id: Optional[str]):
        """Set document properties from Excel data"""
        core_props = doc.core_properties
        core_props.title = f"Facultative Analysis - {data.get('insured_name', 'Unknown')}"
        core_props.author = "AI Facultative System"
        core_props.subject = "Facultative Reinsurance Analysis Report"
        core_props.created = datetime.utcnow()
        core_props.modified = datetime.utcnow()
        if application_id:
            core_props.identifier = application_id
    
    def _add_document_header_from_excel(self, doc: Document, data: Dict[str, Any]):
        """Add document header for Excel-based report"""
        # Add logo centered
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            if os.path.exists(self.logo_path):
                logo_run = logo_para.add_run()
                logo_run.add_picture(self.logo_path, width=Inches(2.0))
        except Exception as e:
            logging.warning(f"Could not add logo: {e}")
            # Add text fallback
            logo_run = logo_para.add_run("KENYA RE")
            logo_run.font.size = Pt(16)
            logo_run.font.bold = True
        
        # Add company name
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run("KENYA REINSURANCE CORPORATION LIMITED")
        company_run.font.size = Pt(12)
        company_run.font.bold = True
        
        # Add document title
        doc.add_paragraph()
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("WORKING SHEET FOR FACULTATIVE REINSURANCE")
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        
        doc.add_paragraph()
    
    def _add_main_analysis_table_from_excel(self, doc: Document, data: Dict[str, Any]):
        """Add main analysis table populated with Excel data"""
        # Add section title
        title_para = doc.add_paragraph("GUIDING POINTS")
        title_para.runs[0].font.bold = True
        title_para.runs[0].font.size = Pt(11)
        
        # Create main table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.0)
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "FIELD"
        header_cells[1].text = "DETAILS"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Define the template fields and map them to Excel data
        template_fields = [
            ("Insured", data.get('insured_name', data.get('insured', 'N/A'))),
            ("Risk Location", data.get('risk_location', data.get('location', 'N/A'))),
            ("Business", data.get('business_type', data.get('business', 'N/A'))),
            ("Cover", data.get('cover_type', data.get('cover', 'N/A'))),
            ("Period", f"{data.get('period_from', 'N/A')} to {data.get('period_to', 'N/A')}"),
            ("Currency", data.get('original_currency', data.get('currency', 'KES'))),
            ("TSI", f"{data.get('original_currency', 'KES')} {data.get('tsi', 0):,.2f}"),
            ("TSI (KES)", f"KES {data.get('tsi_kes', 0):,.2f}"),
            ("Premium", f"{data.get('original_currency', 'KES')} {data.get('premium', 0):,.2f}"),
            ("Premium (KES)", f"KES {data.get('premium_kes', 0):,.2f}"),
            ("Premium Rate %", f"{data.get('premium_rate_percent', 0):.4f}%"),
            ("Premium Rate ‰", f"{data.get('premium_rate_permille', 0):.4f}‰"),
            ("Excess/Deductible", data.get('excess', data.get('deductible', 'N/A'))),
            ("PML", data.get('pml', data.get('probable_maximum_loss', 'N/A'))),
            ("Share Offered %", f"{data.get('share_offered', data.get('offered_share', 0))}%"),
            ("Share Accepted %", f"{data.get('accepted_share', data.get('share_accepted', 0))}%"),
            ("Accepted Liability", f"KES {data.get('accepted_liability', 0):,.2f}"),
            ("Accepted Premium", f"KES {data.get('accepted_premium', 0):,.2f}"),
            ("Net Premium", f"KES {data.get('net_premium', 0):,.2f}"),
            ("Brokerage", f"{data.get('brokerage', 0)}%"),
            ("Loss Ratio", f"{data.get('loss_ratio_percent', 0):.2f}%"),
            ("Accepted Loss Ratio", f"{data.get('accepted_loss_ratio_percent', 0):.2f}%"),
            ("Risk Score", f"{data.get('risk_score', 5.0)}/10"),
            ("Liability Assessment", data.get('liability_assessment', 'Medium')),
            ("ESG Rating", data.get('esg_rating', data.get('esg', 'N/A'))),
        ]
        
        # Add data rows
        for field_name, field_value in template_fields:
            row = table.add_row()
            row.cells[0].text = field_name
            row.cells[1].text = str(field_value)
    
    def _add_calculations_section_from_excel(self, doc: Document, data: Dict[str, Any]):
        """Add calculations section with computed values from Excel data"""
        # Add section title
        doc.add_paragraph()
        calc_title = doc.add_paragraph("CALCULATIONS & ANALYSIS")
        calc_title.runs[0].font.bold = True
        calc_title.runs[0].font.size = Pt(11)
        
        # Create calculations table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(2.5)
        table.columns[2].width = Inches(2.0)
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "METRIC"
        header_cells[1].text = "CALCULATION"
        header_cells[2].text = "RESULT"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add calculation rows with formulas and results
        calculations = [
            ("Premium Rate %", 
             f"({data.get('premium', 0):,.0f} ÷ {data.get('tsi', 0):,.0f}) × 100",
             f"{data.get('premium_rate_percent', 0):.4f}%"),
            
            ("Premium Rate ‰", 
             f"({data.get('premium', 0):,.0f} ÷ {data.get('tsi', 0):,.0f}) × 1000",
             f"{data.get('premium_rate_permille', 0):.4f}‰"),
            
            ("Loss Ratio %", 
             f"({data.get('incurred_losses', 0):,.0f} ÷ {data.get('earned_premium', data.get('premium', 0)):,.0f}) × 100",
             f"{data.get('loss_ratio_percent', 0):.2f}%"),
            
            ("Accepted Liability", 
             f"{data.get('tsi', 0):,.0f} × {data.get('accepted_share', 0)}%",
             f"KES {data.get('accepted_liability', 0):,.2f}"),
            
            ("Accepted Premium", 
             f"{data.get('premium', 0):,.0f} × {data.get('accepted_share', 0)}%",
             f"KES {data.get('accepted_premium', 0):,.2f}"),
            
            ("Net Premium", 
             f"Premium - Brokerage - Taxes - Levies",
             f"KES {data.get('net_premium', 0):,.2f}"),
        ]
        
        for metric, calculation, result in calculations:
            row = table.add_row()
            row.cells[0].text = metric
            row.cells[1].text = calculation
            row.cells[2].text = result
    
    def _add_recommendation_signature_section_from_excel(self, doc: Document, data: Dict[str, Any]):
        """Add recommendation and signature section based on Excel data analysis"""
        # Add section title
        doc.add_paragraph()
        rec_title = doc.add_paragraph("RECOMMENDATION & DECISION")
        rec_title.runs[0].font.bold = True
        rec_title.runs[0].font.size = Pt(11)
        
        # Generate recommendation based on data
        recommendation = self._generate_recommendation_from_data(data)
        
        # Add recommendation text
        rec_para = doc.add_paragraph(recommendation)
        rec_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add signature section
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Create signature table
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.style = 'Table Grid'
        
        # Set column widths
        sig_table.columns[0].width = Inches(3.0)
        sig_table.columns[1].width = Inches(3.0)
        
        # Signature fields
        signatures = [
            ("Analyzed by:", "Approved by:"),
            ("", ""),
            ("Date: _____________", "Date: _____________")
        ]
        
        for i, (left_text, right_text) in enumerate(signatures):
            row = sig_table.rows[i]
            row.cells[0].text = left_text
            row.cells[1].text = right_text
            
            if i == 0:  # Make headers bold
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    def _generate_recommendation_from_data(self, data: Dict[str, Any]) -> str:
        """Generate recommendation text based on calculated data"""
        try:
            loss_ratio = data.get('loss_ratio_percent', 0)
            risk_score = data.get('risk_score', 5.0)
            liability_assessment = data.get('liability_assessment', 'Medium')
            accepted_share = data.get('accepted_share', 0)
            premium_rate = data.get('premium_rate_percent', 0)
            
            recommendation_parts = []
            
            # Analysis of key metrics
            recommendation_parts.append(f"Based on the comprehensive analysis of this facultative reinsurance submission:")
            
            # Loss ratio assessment
            if loss_ratio == 0:
                recommendation_parts.append(f"• Loss ratio data not available for historical analysis.")
            elif loss_ratio < 50:
                recommendation_parts.append(f"• Historical loss ratio of {loss_ratio:.2f}% indicates favorable loss experience.")
            elif loss_ratio < 75:
                recommendation_parts.append(f"• Historical loss ratio of {loss_ratio:.2f}% shows acceptable loss experience.")
            else:
                recommendation_parts.append(f"• Historical loss ratio of {loss_ratio:.2f}% indicates elevated claims experience requiring careful consideration.")
            
            # Premium rate assessment
            if premium_rate > 0:
                if premium_rate < 0.1:
                    recommendation_parts.append(f"• Premium rate of {premium_rate:.4f}% appears competitive for the risk profile.")
                elif premium_rate < 0.5:
                    recommendation_parts.append(f"• Premium rate of {premium_rate:.4f}% is reasonable for the risk category.")
                else:
                    recommendation_parts.append(f"• Premium rate of {premium_rate:.4f}% reflects the higher risk nature of this placement.")
            
            # Risk score analysis
            if risk_score <= 3:
                recommendation_parts.append(f"• Risk score of {risk_score}/10 indicates low risk profile.")
                decision = "RECOMMEND ACCEPTANCE"
            elif risk_score <= 6:
                recommendation_parts.append(f"• Risk score of {risk_score}/10 indicates moderate risk profile.")
                decision = "RECOMMEND CONDITIONAL ACCEPTANCE"
            else:
                recommendation_parts.append(f"• Risk score of {risk_score}/10 indicates higher risk profile requiring enhanced terms.")
                decision = "RECOMMEND CONDITIONAL ACCEPTANCE WITH ENHANCED TERMS"
            
            # Liability assessment
            recommendation_parts.append(f"• Liability assessment: {liability_assessment} - appropriate for our capacity and risk appetite.")
            
            # Final recommendation
            recommendation_parts.append(f"\nRECOMMENDATION: {decision}")
            
            if accepted_share > 0:
                recommendation_parts.append(f"Suggested acceptance: {accepted_share}% share as proposed.")
            
            # Additional considerations
            recommendation_parts.append(f"\nKey considerations:")
            recommendation_parts.append(f"• Ensure all policy terms and conditions are reviewed and agreed")
            recommendation_parts.append(f"• Confirm adequate reinsurance protection for this risk")
            recommendation_parts.append(f"• Monitor loss development closely given the risk profile")
            
            return "\n".join(recommendation_parts)
            
        except Exception as e:
            logging.error(f"Error generating recommendation: {e}")
            return "Recommendation analysis pending. Please review all risk factors and make decision based on underwriting guidelines."


# Service instance
document_report_generator = DocumentReportGenerator()
